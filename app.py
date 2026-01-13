from flask import Flask, render_template, request, send_file, jsonify
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime
import re
from dotenv import load_dotenv
from gigachat import GigaChat

load_dotenv()

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'output'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
app.config['ALLOWED_EXTENSIONS'] = {'docx'}

GIGACHAT_CLIENT_SECRET = os.getenv('GIGACHAT_CLIENT_SECRET')

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    return '\n'.join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

def ask_gigachat(prompt):
    """Запрос к GigaChat"""

    if not GIGACHAT_CLIENT_SECRET:
        return "⚠️ Не настроен GIGACHAT_CLIENT_SECRET"

    try:
        with GigaChat(
            credentials=GIGACHAT_CLIENT_SECRET,
            scope="GIGACHAT_API_PERS",
            verify_ssl_certs=False
        ) as giga:

            full_prompt = f"""Ты - опытный российский юрист, защищающий интересы ПОСТАВЩИКА. 
Отвечай детально и конкретно, только на русском языке.

{prompt}"""

            response = giga.chat(full_prompt)

            if hasattr(response, 'choices') and response.choices:
                return response.choices[0].message.content
            elif hasattr(response, 'content'):
                return response.content
            elif hasattr(response, 'text'):
                return response.text
            else:
                return f"⚠️ Неожиданный формат ответа"

    except Exception as e:
        error_msg = str(e)
        print(f"❌ Ошибка: {error_msg}")

        if "credentials" in error_msg.lower() or "401" in error_msg:
            return "⚠️ Проблема с авторизацией"
        elif "403" in error_msg:
            return "⚠️ Доступ запрещен"
        elif "rate" in error_msg.lower() or "429" in error_msg:
            return "⚠️ Превышен лимит запросов"
        return f"⚠️ Ошибка: {error_msg[:200]}"

def analyze_clause(clause_text, clause_number):
    """Анализ одного пункта договора"""

    prompt = f"""Проанализируй пункт договора с позиции ПОСТАВЩИКА.

ПУНКТ {clause_number}:
{clause_text}

ВАЖНО: Ответь СТРОГО в следующем формате:

РЕДАКЦИЯ_ПОСТАВЩИКА:
[Напиши ПОЛНЫЙ исправленный текст пункта, защищающий интересы Поставщика. 
Используй конкретные цифры: неустойка 0,05-0,1% в день, максимум 5-10%, 
сроки 5-10 рабочих дней, лимит ответственности 30% от договора.
Если исправить невозможно или пункт корректный, напиши: "Редакция Покупателя приемлема"]

КОММЕНТАРИИ:
[Напиши рекомендации для Поставщика:
- В чем риск для Поставщика?
- Почему предложена такая редакция?
- Что важно согласовать дополнительно?
Если редакция приемлема, напиши почему она защищает интересы Поставщика]

Отвечай БЕЗ лишнего текста, строго по формату выше."""

    response = ask_gigachat(prompt)

    if "⚠️" in response:
        return None, response

    # Парсим ответ
    revision = ""
    comments = ""

    if "РЕДАКЦИЯ_ПОСТАВЩИКА:" in response:
        parts = response.split("КОММЕНТАРИИ:")
        revision = parts[0].replace("РЕДАКЦИЯ_ПОСТАВЩИКА:", "").strip()
        if len(parts) > 1:
            comments = parts[1].strip()
    else:
        comments = response

    return revision, comments

def analyze_contract_for_protocol(text):
    """Анализ договора для протокола разногласий"""

    lines = [l.strip() for l in text.split('\n') if l.strip()]

    # Паттерны проблемных мест
    risk_patterns = {
        'высокий': [r'штраф', r'пеня', r'неустойка', r'односторонн', r'без согласования', 
                    r'немедленно', r'не возмещается', r'за счет.*поставщик', r'полная ответственность'],
        'средний': [r'разумный срок', r'своевременно', r'в кратчайшие', r'предоплата', 
                    r'без уведомления', r'по своему усмотрению']
    }

    clauses = []
    clause_number = 1

    for i, line in enumerate(lines, 1):
        line_lower = line.lower()

        # Проверяем, есть ли риск
        has_risk = False
        risk_level = 'низкий'

        for level, patterns in risk_patterns.items():
            for pattern in patterns:
                if re.search(pattern, line_lower):
                    has_risk = True
                    risk_level = level
                    break
            if has_risk:
                break

        # Анализируем только проблемные пункты
        if has_risk and len(line) > 50:  # Достаточно длинный для анализа
            print(f"🤖 Анализирую пункт {clause_number} (строка {i}, риск: {risk_level})...")

            revision, comments = analyze_clause(line, clause_number)

            if revision is not None:
                clauses.append({
                    'number': clause_number,
                    'line': i,
                    'original': line,
                    'revision': revision,
                    'comments': comments,
                    'risk': risk_level
                })
                clause_number += 1

            # Ограничим до 10 пунктов
            if clause_number > 10:
                break

    return clauses

def set_cell_border(cell, **kwargs):
    """Установка границ ячейки таблицы"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '12')
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), '000000')
            tcBorders.append(element)

    tcPr.append(tcBorders)

def create_protocol_word(clauses, filename):
    """Создание Word документа с протоколом разногласий"""

    doc = Document()

    # Настройка полей страницы
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # ЗАГОЛОВОК
    title = doc.add_heading('ПРОТОКОЛ РАЗНОГЛАСИЙ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(16)
    title_run.font.bold = True

    # Подзаголовок
    subtitle = doc.add_paragraph(f'к Договору поставки')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)

    # Дата
    date_para = doc.add_paragraph(f'Дата составления: {datetime.now().strftime("%d.%m.%Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(11)

    doc.add_paragraph()  # Пустая строка

    # Вступительный текст
    intro = doc.add_paragraph(
        'Поставщик предлагает следующие изменения в условия Договора '
        'в целях соблюдения баланса интересов сторон и норм действующего законодательства РФ:'
    )
    intro.runs[0].font.size = Pt(11)

    doc.add_paragraph()  # Пустая строка

    # ТАБЛИЦА
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False

    # Ширина столбцов
    widths = [Inches(2.2), Inches(2.5), Inches(2.5)]
    for i, width in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = width

    # ЗАГОЛОВКИ СТОЛБЦОВ
    header_cells = table.rows[0].cells
    headers = ['РЕДАКЦИЯ ПОКУПАТЕЛЯ', 'РЕДАКЦИЯ ПОСТАВЩИКА', 'КОММЕНТАРИИ']

    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text

        # Форматирование заголовка
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.runs[0]
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(255, 255, 255)

        # Цвет фона заголовка
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '4472C4')  # Синий
        cell._tc.get_or_add_tcPr().append(shading_elm)

        # Границы
        set_cell_border(cell, top={'val': 'single'}, bottom={'val': 'single'}, 
                       left={'val': 'single'}, right={'val': 'single'})

    # СТРОКИ С ДАННЫМИ
    for clause in clauses:
        row_cells = table.add_row().cells

        # Столбец 1: Редакция Покупателя
        cell1 = row_cells[0]
        p1 = cell1.paragraphs[0]

        # Номер пункта жирным
        run_num = p1.add_run(f"Пункт {clause['number']}\n")
        run_num.font.bold = True
        run_num.font.size = Pt(10)

        # Оригинальный текст
        run_text = p1.add_run(clause['original'])
        run_text.font.size = Pt(10)

        # Столбец 2: Редакция Поставщика
        cell2 = row_cells[1]
        p2 = cell2.paragraphs[0]
        run2 = p2.add_run(clause['revision'])
        run2.font.size = Pt(10)

        # Если редакция приемлема - зеленым
        if "приемлема" in clause['revision'].lower() or "принимается" in clause['revision'].lower():
            run2.font.color.rgb = RGBColor(0, 128, 0)
            run2.font.bold = True

        # Столбец 3: Комментарии
        cell3 = row_cells[2]
        p3 = cell3.paragraphs[0]
        run3 = p3.add_run(clause['comments'])
        run3.font.size = Pt(9)
        run3.font.italic = True

        # Границы ячеек
        for cell in row_cells:
            set_cell_border(cell, top={'val': 'single'}, bottom={'val': 'single'}, 
                           left={'val': 'single'}, right={'val': 'single'})

            # Отступы в ячейках
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(4)
                paragraph.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()  # Пустая строка

    # ОБОСНОВАНИЕ
    doc.add_heading('Правовое обоснование', level=2)

    justification = doc.add_paragraph()
    justification.add_run(
        'Предложенные изменения соответствуют нормам Гражданского кодекса РФ:\n'
    ).font.size = Pt(10)

    articles = [
        '• Статья 330-333 ГК РФ — о неустойке и её соразмерности',
        '• Статья 421 ГК РФ — о свободе договора',
        '• Статья 422 ГК РФ — о соответствии договора закону',
        '• Статья 450-453 ГК РФ — об изменении и расторжении договора'
    ]

    for article in articles:
        p = doc.add_paragraph(article, style='List Bullet')
        p.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # РЕКОМЕНДАЦИИ
    doc.add_heading('Стандартные параметры по законодательству РФ', level=2)

    standards = doc.add_paragraph()
    standards_text = (
        '📊 Неустойка за просрочку: 0,05-0,1% в день, но не более 5-10% от суммы обязательства\n'
        '⏰ Сроки уведомлений: 3-5 рабочих дней\n'
        '⏰ Сроки устранения недостатков: 5-10 рабочих дней\n'
        '💰 Общее ограничение ответственности: не более 30% от суммы договора\n'
        '⚖️ Основание: практика арбитражных судов РФ и статья 333 ГК РФ о снижении неустойки'
    )
    standards.add_run(standards_text).font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph()

    # ПОДПИСИ
    signature_table = doc.add_table(rows=3, cols=2)
    signature_table.autofit = False

    # Заголовки
    signature_table.cell(0, 0).text = 'ПОСТАВЩИК'
    signature_table.cell(0, 1).text = 'ПОКУПАТЕЛЬ'

    for i in range(2):
        cell = signature_table.rows[0].cells[i]
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)

    # Поля для подписей
    signature_table.cell(1, 0).text = '_' * 30
    signature_table.cell(1, 1).text = '_' * 30

    signature_table.cell(2, 0).text = '(подпись, печать)'
    signature_table.cell(2, 1).text = '(подпись, печать)'

    for i in range(2):
        signature_table.rows[2].cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        signature_table.rows[2].cells[i].paragraphs[0].runs[0].font.italic = True

    doc.add_paragraph()

    # ФУТЕР
    footer = doc.add_paragraph(
        f'Протокол разногласий составлен {datetime.now().strftime("%d.%m.%Y")} '
        f'с использованием GigaChat (Сбербанк) 🇷🇺'
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.italic = True
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # Сохранение
    path = os.path.join(app.config['OUTPUT_FOLDER'], filename)
    doc.save(path)
    return path

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload():
    if 'contract' not in request.files:
        return jsonify({'error': 'Нет файла'}), 400

    file = request.files['contract']
    if not file or file.filename == '':
        return jsonify({'error': 'Файл не выбран'}), 400

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(path)

        try:
            text = extract_text_from_docx(path)
            if len(text) < 50:
                return jsonify({'error': 'Документ слишком короткий'}), 400

            print(f"📄 Документ загружен: {len(text)} символов")
            print("🤖 Анализирую договор с позиции Поставщика...")

            clauses = analyze_contract_for_protocol(text)

            if not clauses:
                return jsonify({'error': 'Не найдено проблемных пунктов для анализа'}), 400

            print(f"✅ Проанализировано {len(clauses)} пунктов")
            print("📝 Создаю протокол разногласий...")

            out = f'Протокол_разногласий_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            create_protocol_word(clauses, out)

            # Формируем текстовый вывод для отображения
            analysis_text = f"""
{'='*80}
ПРОТОКОЛ РАЗНОГЛАСИЙ К ДОГОВОРУ ПОСТАВКИ
{'='*80}

Дата: {datetime.now().strftime("%d.%m.%Y %H:%M")}
Проанализировано пунктов: {len(clauses)}

"""

            for clause in clauses:
                analysis_text += f"""
{'─'*80}
ПУНКТ {clause['number']} (строка {clause['line']})
{'─'*80}

📄 РЕДАКЦИЯ ПОКУПАТЕЛЯ:
{clause['original']}

✏️ РЕДАКЦИЯ ПОСТАВЩИКА:
{clause['revision']}

💬 КОММЕНТАРИИ:
{clause['comments']}

"""

            analysis_text += f"""
{'='*80}
⚖️ ПРАВОВОЕ ОБОСНОВАНИЕ
{'='*80}

Предложенные изменения соответствуют:
• ГК РФ статьи 330-333 (неустойка)
• ГК РФ статьи 421-422 (свобода договора)
• ГК РФ статьи 450-453 (изменение договора)

📊 СТАНДАРТНЫЕ ПАРАМЕТРЫ:
• Неустойка: 0,05-0,1% в день, макс 5-10%
• Сроки уведомлений: 3-5 рабочих дней
• Сроки устранения: 5-10 рабочих дней
• Лимит ответственности: 30% от договора

🇷🇺 Анализ выполнен GigaChat (Сбербанк)
{'='*80}
"""

            print("✅ Протокол разногласий готов!")

            return jsonify({
                'success': True,
                'analysis': analysis_text,
                'protocol': analysis_text,
                'download_url': f'/download/{out}',
                'risks_found': len(clauses)
            })

        except Exception as e:
            import traceback
            traceback.print_exc()
            return jsonify({'error': f'Ошибка: {str(e)}'}), 500

    return jsonify({'error': 'Только .docx файлы'}), 400

@app.route('/download/<filename>')
def download(filename):
    return send_file(
        os.path.join(app.config['OUTPUT_FOLDER'], filename),
        as_attachment=True,
        download_name=filename
    )

if __name__ == '__main__':
    print("="*60)
    print("🚀 AI CONTRACT ANALYZER - ПРОТОКОЛ РАЗНОГЛАСИЙ")
    print("🤖 GigaChat (Сбербанк)")
    print("🇷🇺 Российский ИИ")
    print("⚖️ Защита интересов ПОСТАВЩИКА")
    print("="*60)

    if not GIGACHAT_CLIENT_SECRET:
        print("❌ GIGACHAT_CLIENT_SECRET не найден!")
    else:
        print(f"✅ Client Secret: {GIGACHAT_CLIENT_SECRET[:10]}...")
        print()
        print("🧪 Тест подключения...")
        test = ask_gigachat("Ответь коротко: работаешь?")
        if "⚠️" in test:
            print(f"❌ Тест не прошел:\n{test}")
        else:
            print(f"✅ Тест прошел! {test[:50]}...")

    print()
    print("🌐 http://localhost:5000")
    print("="*60)

    app.run(debug=True, host='0.0.0.0', port=5000)
